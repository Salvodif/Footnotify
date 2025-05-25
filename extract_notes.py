import os
from docx import Document as DocxDocument
from odf.opendocument import OpenDocumentText, load as odf_load
from odf.text import P as OdfP, Span as OdfSpan, Note as OdfNote, NoteBody as OdfNoteBody, LineBreak as OdfLineBreak
from odf.style import Style, TextProperties
from odf import teletype
from odf.element import Element 
from odf.namespaces import OFFICENS

import settings_parser
import footnote_processor


# Importazioni Rich
from rich.console import Console
from rich.panel import Panel
from rich.prompt import Prompt # Per un input più carino con Rich
from rich.progress import Progress, SpinnerColumn, BarColumn, TextColumn, TimeElapsedColumn
from rich.rule import Rule

# Inizializza la console Rich
console = Console()

# --- Funzioni Helper per Stili ODT --- (Identiche a prima)
def get_odt_style_name(fmt_dict):
    parts = []
    if fmt_dict.get('bold'): parts.append("bold")
    if fmt_dict.get('italic'): parts.append("italic")
    if fmt_dict.get('underline'): parts.append("underline")
    
    color = fmt_dict.get('color')
    if color:
        # Sanitize color for style name: remove '#' and use 'c' prefix
        sanitized_color = color[1:] if color.startswith('#') else color
        parts.append(f"c{sanitized_color}")

    if not parts: return "default" # Default style name if no formatting
    return "ft_" + "_".join(parts)

def create_or_get_odt_style(doc_automatic_styles_container, fmt_dict):
    style_name = get_odt_style_name(fmt_dict)
    if style_name == "default": 
        # For default, we don't apply a specific named style, 
        # letting it inherit paragraph default or application default.
        return None

    # Check if style already exists in the provided container
    if doc_automatic_styles_container:
        for s_element in doc_automatic_styles_container.getElementsByType(Style):
            # Ensure it's a text style and the name matches
            if s_element.getAttribute("name") == style_name and \
               s_element.getAttribute("family") == "text":
                return style_name # Return existing style name
            
    # Create new style if it doesn't exist
    style = Style(name=style_name, family="text")
    text_props = TextProperties()
    
    has_any_property = False # Track if any text properties are set

    # Standard font properties
    if fmt_dict.get('bold'):
        text_props.setAttribute("fontweight", "bold")
        has_any_property = True
    if fmt_dict.get('italic'):
        text_props.setAttribute("fontstyle", "italic")
        has_any_property = True
    if fmt_dict.get('underline'):
        # ODF standard properties for simple underline
        text_props.setAttribute("textunderlinetype", "single")
        text_props.setAttribute("textunderlinestyle", "solid")
        # text_props.setAttribute("textunderlinecolor", "font-color") # Inherits current font color
        has_any_property = True
    
    # Color property
    color_value = fmt_dict.get('color')
    if color_value:
        text_props.setAttribute("color", color_value) # ODF uses 'color' for font color
        has_any_property = True

    if has_any_property:
        style.addElement(text_props)
        if doc_automatic_styles_container:
            doc_automatic_styles_container.addElement(style)
        else:
            # This case should ideally not happen if automaticstyles is always present
            console.print(f"[yellow]AVVISO:[/yellow] Contenitore stili automatici non valido o non fornito per lo stile '{style_name}'. Lo stile potrebbe non essere salvato correttamente.")
        return style_name # Return the name of the created or found style
    
    # If no properties were applied (e.g., empty fmt_dict or only non-style properties),
    # it's effectively a default style, so no specific named style is needed.
    return None

# --- Funzioni di Estrazione Footnote --- (Identiche a prima)
def extract_footnotes_from_docx(filepath):
    console.print(f"Estrazione footnotes da DOCX: [cyan]{filepath}[/cyan]")
    doc = DocxDocument(filepath)
    extracted_footnotes = []

    footnotes_part = None
    if hasattr(doc.part, 'document') and hasattr(doc.part.document, 'footnotes_part'):
        footnotes_part = doc.part.document.footnotes_part
    
    if not footnotes_part or not footnotes_part.footnotes:
        console.print("[yellow]Nessuna footnote trovata nel documento DOCX.[/yellow]")
        return []

    num_footnotes = len(footnotes_part.footnotes)
    with Progress(
        SpinnerColumn(),
        TextColumn("[progress.description]{task.description}"),
        BarColumn(),
        TextColumn("[progress.percentage]{task.percentage:>3.0f}%"),
        TimeElapsedColumn(),
        console=console,
        transient=True
    ) as progress:
        task = progress.add_task(f"Elaborazione {num_footnotes} footnotes DOCX...", total=num_footnotes)
        for i, footnote in enumerate(footnotes_part.footnotes):
            footnote_paragraphs_data = []
            for paragraph in footnote.paragraphs:
                paragraph_runs_data = []
                for run in paragraph.runs:
                    fmt = {
                        'bold': run.bold if run.bold is not None else False,
                        'italic': run.italic if run.italic is not None else False,
                        'underline': run.underline if run.underline is not None else False,
                    }
                    paragraph_runs_data.append((run.text, fmt))
                footnote_paragraphs_data.append(paragraph_runs_data)
            extracted_footnotes.append(footnote_paragraphs_data)
            progress.update(task, advance=1)
    
    console.print(f"[green]Estratte {len(extracted_footnotes)} footnotes da DOCX.[/green]")
    return extracted_footnotes

def parse_odt_style_attributes(doc_odt, style_name_to_find):
    fmt = {'bold': False, 'italic': False, 'underline': False}
    if not style_name_to_find:
        return fmt

    style_element = None
    style_sources = []
    if hasattr(doc_odt, 'automaticstyles') and doc_odt.automaticstyles:
        style_sources.append(doc_odt.automaticstyles)
    if hasattr(doc_odt, 'styles') and doc_odt.styles:
        style_sources.append(doc_odt.styles)
    
    MODEL_STYLE_QNAME = Style(name="temp-qname-style", family="text").qname

    for source_container in style_sources:
        for s_candidate in source_container.childNodes:
            if hasattr(s_candidate, 'qname') and s_candidate.qname == MODEL_STYLE_QNAME:
                 if s_candidate.getAttribute("name") == style_name_to_find and \
                    s_candidate.getAttribute("family") == "text":
                    style_element = s_candidate
                    break
            if style_element: break
        if style_element: break
    
    if style_element:
        text_props_list = style_element.getElementsByType(TextProperties)
        if text_props_list:
            text_props = text_props_list[0]
            if text_props.getAttribute("fontweight") == "bold":
                fmt['bold'] = True
            if text_props.getAttribute("fontstyle") == "italic":
                fmt['italic'] = True
            if text_props.getAttribute("textunderlinetype") and \
               text_props.getAttribute("textunderlinetype") != "none":
                fmt['underline'] = True
    return fmt


def extract_footnotes_from_odt(filepath):
    console.print(f"Estrazione footnotes da ODT: [cyan]{filepath}[/cyan]")
    doc = odf_load(filepath)
    extracted_footnotes = []

    notes_found = doc.getElementsByType(OdfNote)
    
    MODEL_NOTE_QNAME = OdfNote(noteclass="footnote", id="temp-qname").qname
    MODEL_NOTE_BODY_QNAME = OdfNoteBody().qname
    MODEL_P_QNAME = OdfP().qname
    MODEL_SPAN_QNAME = OdfSpan().qname
    MODEL_LINEBREAK_QNAME = OdfLineBreak().qname

    footnote_elements = []
    for n in notes_found:
        if hasattr(n, 'qname') and n.qname == MODEL_NOTE_QNAME: # Confronto qname
            if n.getAttribute("noteclass") == "footnote": # Poi attributo
                footnote_elements.append(n)

    if not footnote_elements:
        console.print("[yellow]Nessuna footnote trovata nel documento ODT.[/yellow]")
        return []

    num_footnotes = len(footnote_elements)
    with Progress(
        SpinnerColumn(),
        TextColumn("[progress.description]{task.description}"),
        BarColumn(),
        TextColumn("[progress.percentage]{task.percentage:>3.0f}%"),
        TimeElapsedColumn(),
        console=console,
        transient=True
    ) as progress:
        task = progress.add_task(f"Elaborazione {num_footnotes} footnotes ODT...", total=num_footnotes)
        for i, note_element in enumerate(footnote_elements):
            note_body = None
            for child in note_element.childNodes:
                if hasattr(child, 'qname') and child.qname == MODEL_NOTE_BODY_QNAME:
                    note_body = child
                    break
            
            if not note_body:
                progress.update(task, advance=1)
                continue
            
            footnote_paragraphs_data = []
            for p_element_candidate in note_body.childNodes:
                if not (hasattr(p_element_candidate, 'qname') and p_element_candidate.qname == MODEL_P_QNAME):
                    continue

                p_element = p_element_candidate
                paragraph_runs_data = []
                
                base_fmt = {} 
                # parent_p_style_name = p_element.getAttribute("stylename")
                # if parent_p_style_name:
                #     base_fmt = parse_odt_style_attributes(doc, parent_p_style_name)


                for child_node in p_element.childNodes:
                    text_content = "" 
                    fmt = base_fmt.copy() 

                    if child_node.nodeType == child_node.TEXT_NODE:
                        text_content = child_node.data
                    elif hasattr(child_node, 'qname'):
                        if child_node.qname == MODEL_SPAN_QNAME:
                            span_element = child_node
                            text_content = teletype.extractText(span_element)
                            style_name = span_element.getAttribute("stylename")
                            if style_name:
                                span_fmt = parse_odt_style_attributes(doc, style_name)
                                fmt.update(span_fmt)
                        elif child_node.qname == MODEL_LINEBREAK_QNAME:
                            text_content = "\n"
                    
                    if text_content.strip() or text_content == "\n":
                        paragraph_runs_data.append((text_content, fmt))
                
                if paragraph_runs_data:
                     footnote_paragraphs_data.append(paragraph_runs_data)
            extracted_footnotes.append(footnote_paragraphs_data)
            progress.update(task, advance=1)

    console.print(f"[green]Estratte {len(extracted_footnotes)} footnotes da ODT.[/green]")
    return extracted_footnotes


# --- Funzione di Creazione ODT con Footnotes --- (Identica a prima)
def create_odt_with_footnotes(footnotes_data, output_filepath):
    console.print(f"Creazione file ODT di output: [cyan]{output_filepath}[/cyan]")
    new_doc = OpenDocumentText()

    if not hasattr(new_doc, 'automaticstyles') or new_doc.automaticstyles is None:
        console.print("[yellow]INFO:[/yellow] [cyan]new_doc.automaticstyles[/cyan] non inizializzato. Tentativo di creazione manuale.")
        auto_styles_element = Element(qname=(OFFICENS, 'automatic-styles'))
        if hasattr(new_doc, 'document') and new_doc.document is not None:
            target_parent = new_doc.document
            reference_node = None
            if hasattr(new_doc, 'body') and new_doc.body is not None:
                reference_node = new_doc.body
            if reference_node:
                target_parent.insertBefore(auto_styles_element, reference_node)
            else:
                if target_parent.hasChildNodes():
                    target_parent.insertBefore(auto_styles_element, target_parent.firstChild)
                else:
                    target_parent.addElement(auto_styles_element)
            new_doc.automaticstyles = auto_styles_element
            console.print("[green]INFO:[/green] Creato e inserito [cyan]<office:automatic-styles>[/cyan] manualmente.")
        else:
            console.print("[bold red]ERRORE:[/bold red] Impossibile trovare [cyan]new_doc.document[/cyan] per inserire [cyan]automatic-styles[/cyan].")

    num_footnotes_to_write = len(footnotes_data)
    with Progress(
        SpinnerColumn(),
        TextColumn("[progress.description]{task.description}"),
        BarColumn(),
        TextColumn("[progress.percentage]{task.percentage:>3.0f}%"),
        TimeElapsedColumn(),
        console=console,
        transient=True
    ) as progress:
        task_write = progress.add_task(f"Scrittura {num_footnotes_to_write} footnotes...", total=num_footnotes_to_write)
        for i, footnote_content_paragraphs in enumerate(footnotes_data):
            heading_p = OdfP()
            heading_span = OdfSpan(text=f"--- Nota a piè di pagina {i+1} ---")
            heading_p.addElement(heading_span)
            new_doc.text.addElement(heading_p)

            for paragraph_runs in footnote_content_paragraphs:
                if not paragraph_runs:
                    new_doc.text.addElement(OdfP())
                    continue
                odt_paragraph = OdfP()
                for text_content, fmt_dict in paragraph_runs:
                    if not text_content: continue
                    if "\n" in text_content:
                        parts = text_content.split('\n')
                        first_part = True
                        for part_text in parts:
                            if not first_part and odt_paragraph.hasChildNodes():
                                odt_paragraph.addElement(OdfLineBreak())
                            if part_text:
                                span = OdfSpan(text=part_text)
                                style_name = create_or_get_odt_style(new_doc.automaticstyles, fmt_dict)
                                if style_name:
                                    span.setAttribute("stylename", style_name)
                                odt_paragraph.addElement(span)
                            first_part = False
                    else:
                        span = OdfSpan(text=text_content)
                        style_name = create_or_get_odt_style(new_doc.automaticstyles, fmt_dict)
                        if style_name:
                            span.setAttribute("stylename", style_name)
                        odt_paragraph.addElement(span)
                if odt_paragraph.hasChildNodes():
                    new_doc.text.addElement(odt_paragraph)
            new_doc.text.addElement(OdfP())
            progress.update(task_write, advance=1)
    try:
        new_doc.save(output_filepath)
        console.print(f"[green]File '{output_filepath}' creato con successo con {num_footnotes_to_write} footnotes.[/green]")
    except Exception as e:
        console.print(f"[bold red]Errore durante il salvataggio del file ODT:[/bold red] {e}")
        # (gestione errore directory omessa per brevità, ma era corretta prima)

# --- Funzione Principale Modificata ---
def run_footnote_extraction(input_filepath, output_filepath_base="output_footnotes", settings_filepath="settings.yaml"):
    """
    Funzione principale che orchestra l'estrazione e la creazione del file di output.
    Permette di specificare un file di settings alternativo.
    """
    if not os.path.exists(input_filepath):
        console.print(f"[bold red]Errore:[/bold red] File di input '[cyan]{input_filepath}[/cyan]' non trovato.")
        return

    # Crea la directory di output se non esiste
    output_dir = os.path.abspath(output_filepath_base) # Ensure output_dir is absolute for clarity
    if not os.path.exists(output_dir):
        try:
            os.makedirs(output_dir)
            console.print(f"Creata directory di output: [cyan]{output_dir}[/cyan]")
        except OSError as e:
            console.print(f"[bold red]Errore creando directory di output {output_dir}:[/bold red] {e}")
            return

    base_filename, ext = os.path.splitext(os.path.basename(input_filepath))
    ext = ext.lower()

    output_filename = f"footnotes_da_{base_filename.replace(' ', '_')}.odt"
    output_filepath_full = os.path.join(output_dir, output_filename)
    
    console.print(Rule(f"Inizio elaborazione: [b cyan]{os.path.basename(input_filepath)}[/b cyan]"))
    console.print(f"File di settings in uso: [magenta]{os.path.abspath(settings_filepath)}[/magenta]")
    
    settings = settings_parser.load_settings(settings_filepath)
    if not settings:
        console.print(f"[bold red]Errore: Impossibile caricare il file delle impostazioni '{os.path.abspath(settings_filepath)}'. Elaborazione annullata.[/bold red]")
        return
    
    raw_extracted_footnotes_data = []
    if ext == ".docx":
        try:
            doc_for_check = DocxDocument(input_filepath)
            fp = None
            if hasattr(doc_for_check.part, 'document') and hasattr(doc_for_check.part.document, 'footnotes_part'):
                 fp = doc_for_check.part.document.footnotes_part
            
            if fp and fp.footnotes:
                 raw_extracted_footnotes_data = extract_footnotes_from_docx(input_filepath)
            else:
                console.print("[yellow]Nessuna footnote trovata nel documento DOCX (controllo preliminare).[/yellow]")
        except Exception as e:
            console.print(f"[bold red]Errore durante il caricamento o controllo preliminare del DOCX:[/bold red] {e}")
            return

    elif ext == ".odt":
        raw_extracted_footnotes_data = extract_footnotes_from_odt(input_filepath)
    else:
        console.print(f"[bold red]Errore:[/bold red] Formato file '[yellow]{ext}[/yellow]' non supportato. Usare .docx o .odt.")
        return

    if raw_extracted_footnotes_data:
        console.print(f"Numero di footnote estratte raw: {len(raw_extracted_footnotes_data)}")
        console.print("Inizio processamento e formattazione footnotes...")
        
        processed_footnotes_for_odt = []
        with Progress(
            SpinnerColumn(),
            TextColumn("[progress.description]{task.description}"),
            BarColumn(),
            TextColumn("[progress.percentage]{task.percentage:>3.0f}%"),
            TimeElapsedColumn(),
            console=console,
            transient=True
        ) as progress:
            task_process = progress.add_task(f"Processando {len(raw_extracted_footnotes_data)} footnotes...", total=len(raw_extracted_footnotes_data))

            for single_raw_footnote_data in raw_extracted_footnotes_data:
                match_results = footnote_processor.parse_and_match_footnote(
                    single_raw_footnote_data, settings
                )

                # Ensure parsed_fields has preprocessed_text if no match, for format_parsed_footnote
                # This was handled in footnote_processor.py:
                # if match_results['matched_type'] is None:
                #    match_results['parsed_fields']['preprocessed_text'] = match_results['preprocessed_text']

                formatted_segments = footnote_processor.get_formatted_and_colored_footnote_segments(
                    match_results['matched_type'],
                    match_results['parsed_fields'], # This will contain 'preprocessed_text' if no match
                    settings,
                    match_results['confidence']
                )
                
                # create_odt_with_footnotes expects a list of footnotes,
                # where each footnote is a list of paragraphs,
                # and each paragraph is a list of (text, fmt) tuples.
                # formatted_segments is already a list of (text, fmt) tuples, representing one paragraph.
                if formatted_segments: 
                    processed_footnotes_for_odt.append([formatted_segments]) 
                else:
                    # This case should ideally be less common if unmatched footnotes are formatted with preprocessed_text
                    # If format_parsed_footnote returns empty for an unmatched note (e.g. preprocessed_text was also empty)
                    # we add an empty paragraph for that footnote to maintain footnote count.
                    processed_footnotes_for_odt.append([[]]) 
                progress.update(task_process, advance=1)

        console.print(f"Processate {len(processed_footnotes_for_odt)} footnotes per l'output.")
        create_odt_with_footnotes(processed_footnotes_for_odt, output_filepath_full)
    else:
        console.print("[yellow]Nessuna footnote da scrivere nel file di output.[/yellow]")
    
    console.print(Rule(f"Fine elaborazione: [b cyan]{os.path.basename(input_filepath)}[/b cyan]"))


# --- Esecuzione Script ---

def run_tests(test_files_dir="test_data"):
    """
    Esegue una serie di test predefiniti utilizzando file di test.
    """
    console.print(Rule("[bold yellow]MODALITÀ TEST[/bold yellow]", style="yellow"))
    
    # Ensure test_files_dir is absolute or relative to the script's location if needed
    # For simplicity, assuming test_data is in the same directory as the script or CWD
    abs_test_files_dir = os.path.abspath(test_files_dir)
    console.print(f"Directory dei file di test: [cyan]{abs_test_files_dir}[/cyan]")

    test_settings = os.path.join(abs_test_files_dir, "test_settings.yaml")
    if not os.path.exists(test_settings):
        console.print(f"[bold red]Errore File Test:[/bold red] Il file di settings per i test '{test_settings}' non è stato trovato.")
        console.print(f"Assicurati che '{os.path.join(test_files_dir, 'test_settings.yaml')}' esista e sia configurato.")
        return

    test_files = {
        "DOCX Test": os.path.join(abs_test_files_dir, "test_document.docx"),
        "ODT Test": os.path.join(abs_test_files_dir, "test_document.odt"),
    }

    output_base_test_dir = "output_tests" # Base directory for all test outputs

    for test_name, test_doc_path in test_files.items():
        console.print(f"\n--- Esecuzione Test: [b cyan]{test_name}[/b cyan] ---")
        if not os.path.exists(test_doc_path):
            console.print(f"[bold red]Errore File Test:[/bold red] Documento di test '{test_doc_path}' non trovato.")
            console.print(f"Per favore, crea il file '{os.path.basename(test_doc_path)}' nella directory '{abs_test_files_dir}'.")
            # Detailed instructions for creating the files are good here, as per the task.
            if "DOCX" in test_name:
                console.print("Contenuto atteso per DOCX ('test_document.docx'):")
                console.print("  - Main: \"This is a test document for DOCX. Book: Green.^(F1) Journal: Yellow.^(F2) No Match: Red.^(F3) Special Classic: Green.^(F4)\"")
                console.print("  - F1 (Book Green): \"John Author, <i>The Great Book</i> (Cityville: PubCo, 2023). Edition 2.\" (Title italic)")
                console.print("  - F2 (Journal Yellow): \"Jane Writer, \"An Article Title,\" <i>Journal of Studies</i> (2022).\" (Title in quotes, Journal italic)")
                console.print("  - F3 (No Match Red): \"Unmatchable Reference, Some Details, 1999.\"")
                console.print("  - F4 (Special Classic Green): \"STh II-II, q. 47, a. 1.\"")
            elif "ODT" in test_name:
                console.print("Contenuto atteso per ODT ('test_document.odt'):")
                console.print("  - Main: \"This is an ODT test document. Book ODT: Green.^(F1) Web ODT: Yellow.^(F2) Minimal ODT: Red.^(F3) Special Classic ODT: Green.^(F4)\"")
                console.print("  - F1 (Book Green): \"Alice Reader, <i>Another Fine Tome</i> (Townsfolk: PrintPress, 2021). ISBN 12345. Edition 1.\" (Title italic)")
                console.print("  - F2 (Web Yellow): \"Bob Online, \"Web Thoughts,\" <i>Bob's Site</i>.\" (Title in quotes, Site italic)")
                console.print("  - F3 (No Match Red): \"Just a few words.\"")
                console.print("  - F4 (Special Classic Green): \"DS 4501.\"")
            console.print(f"Skipping test: {test_name}")
            continue

        # Create a specific output subdirectory for this test run's outputs
        output_specific_dir = os.path.join(output_base_test_dir, f"output_{test_name.lower().replace(' ', '_')}")
        
        run_footnote_extraction(
            input_filepath=test_doc_path,
            output_filepath_base=output_specific_dir, # Pass the specific dir for this test
            settings_filepath=test_settings
        )
        
        # Construct the expected output file name to inform the user
        base_input_filename = os.path.splitext(os.path.basename(test_doc_path))[0]
        expected_output_filename = f"footnotes_da_{base_input_filename.replace(' ', '_')}.odt"
        full_output_path = os.path.join(output_specific_dir, expected_output_filename)

        console.print(f"[green]Test '{test_name}' completato.[/green] Output (potenzialmente) generato in: [b magenta]{os.path.abspath(full_output_path)}[/b magenta]")
        console.print("--> [bold]VERIFICA MANUALE RICHIESTA:[/bold] Apri il file e controlla la formattazione e i colori delle note.")

    console.print(Rule("Fine Modalità Test", style="yellow"))


if __name__ == "__main__":
    console.print(Panel.fit("📝 Estrattore di Footnote Interattivo 📝", style="bold blue on white", border_style="blue"))
    
    while True:
        console.print(Rule(style="dim blue")) # Visual separator for clarity
        input_choice = Prompt.ask(
            "[b]Inserisci il percorso del file .odt/.docx[/b], '[b]test[/b]' per la modalità test, o '[b]esci[/b]' per terminare"
        ).strip()

        if not input_choice:
            console.print("[yellow]Nessun input. Riprova o digita 'esci' o 'test'.[/yellow]")
            continue

        if input_choice.lower() == 'esci':
            console.print("[bold blue]Arrivederci![/bold blue]")
            break
        
        if input_choice.lower() == 'test':
            run_tests()
            continue

        if not os.path.exists(input_choice):
            console.print(f"[bold red]Errore:[/bold red] Il file '[cyan]{input_choice}[/cyan]' non esiste. Controlla il percorso e riprova.")
            continue

        _, ext_check = os.path.splitext(input_choice.lower())
        if ext_check not in [".odt", ".docx"]:
            console.print(f"[bold red]Errore:[/bold red] Estensione file '[yellow]{ext_check}[/yellow]' non supportata. Sono accettati solo .odt e .docx.")
            continue
            
        # For regular operation, use default settings.yaml and standard output directory
        run_footnote_extraction(input_choice, output_filepath_base="output_footnotes") 
        console.print("\nElaborazione completata per questo file.")

    console.print("[bold green]Programma terminato.[/bold green]")